import os
import json
import logging
import re
from datetime import datetime
from dotenv import load_dotenv
import pandas as pd
import openai
from docx import Document
from sharepoint_fields.requestSharepointFields import requestSharepointFields

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler(), logging.FileHandler("pipeline_download.log")]
)

# Load environment variables from .env
load_dotenv()
CLIENT_ID = os.getenv("CLIENT_ID")
CLIENT_SECRET = os.getenv("CLIENT_SECRET")
TENANT_ID = os.getenv("TENANT_ID")
openai.api_key = os.getenv("OPENAI_API_KEY")
OpenAIclient = openai.OpenAI(api_key=openai.api_key)
A2_PIPELINE = os.getenv("A2_PIPELINE")
A2_CV_FOLDER = os.getenv("A2_CV_FOLDER")
CONSULTANT_JSON_FILE = "consultant_cv_data.json"
PROFILES_FOLDER = "consultant_profiles"
os.makedirs(PROFILES_FOLDER, exist_ok=True)  # Ensure profiles folder exists

CompetencyBotIntro = """
Du er en tjeneste som mottar teksten fra en CV og skal returnere informasjon i en presis, flat JSON-struktur. Følg instruksjonene nøye og gi konsistente svar.
Generelle regler:
1. Bruk eksakt samme nøkkelnavn som angitt her. Ingen variasjoner i store og små bokstaver, ingen ekstra mellomrom eller spesialtegn.
2. Returner alle felter, selv om verdien er ukjent. Hvis data mangler, bruk verdien null.
3. Bruk kun tall uten enhet (f.eks. 5, ikke 5 måneder) for felt som krever et antall.
4. Returner alle svar som en flat JSON-struktur uten underliggende objekter eller lister.
5. For felter som ender med .mnd skal tallet være antall måneder konsulenten har jobbet med dette eller 0 hvis aldri.
6. For felter som ender med .sist skal tallet være antall måneder siden konsulenten sist jobbet med dette, MEN settes til null hvis .mnd feltet =0.

Feltstruktur: Besvar med følgende JSON-struktur:
"""

CompetencyBotOutro = """
Viktige merknader:
- For hver område, bruk nøyaktig samme format for spesifiserte felter med .mnd, .sist, og .ref som angitt.
- Eksempel: Hvis konsulenten ikke har erfaring i bransjen Forsvar, sett "Forsvar.mnd": 0, "Forsvar.sist": null, og "Forsvar.ref": "".
"""

CompetencyBotSchemas = ["Schema_bransjeerfaring.json", "Schema_teknologi_kompetanse.json", "Schema_annen_kompetanse.json"]

# Function to extract text from .docx files
def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        logging.error(f"Error reading DOCX file {docx_path}: {e}")
        return ""

# Function to anonymize consultant's name in the CV text
def anonymizeName(file_body, consultant_name, consultant_alias="Konsulenten"):
    name_split = consultant_name.split()
    first_name, last_name = name_split[0], name_split[-1]
    short_name = first_name

    content = file_body
    content = re.sub(r"\b%s\b" % consultant_name, f"[{consultant_alias}]", content)
    content = re.sub(r"\b%s\b" % first_name, consultant_alias, content)
    content = re.sub(r"\b%s\b" % last_name, consultant_alias, content)
    content = re.sub(r"\b%s\b" % short_name, consultant_alias, content)

    return content

def Query_OpenAI_Completion(instruction, query_text):
    # Call the OpenAI API for a general ChatGPT chat completion
    try:
        # Call OpenAI's completion API for ChatGPT
        completion = OpenAIclient.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": instruction},
                {
                    "role": "user",
                    "content": query_text
                }
            ]
        )
        # Extract and return the OpenAI ChatgTP response
        #return completion.choices[0].message
        # Extract and return the OpenAI ChatGPT response content only
        response_content = completion.choices[0].message.content
        
        # Check if the response content is in JSON format and strip it if needed
        if response_content.startswith("```json"):
            response_content = response_content.strip("```json\n").strip("```")
        
        return response_content
    
    except Exception as e:
        logging.warning(f"An error occurred: {e}")
        return None

import json

def Query_OpenAI_Completion_JSON(instruction, query_text):
    """
    Calls the OpenAI API to analyze CV content and return the result as a JSON object.
    
    Args:
        instruction (str): The system instruction for the assistant.
        query_text (str): The consultant's anonymized CV text.

    Returns:
        dict: Parsed JSON response from OpenAI if successful, or None if there's an error.
    """
    try:
        # Call OpenAI's completion API for ChatGPT
        completion = OpenAIclient.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": instruction},
                {"role": "user", "content": query_text}
            ]
        )
        
        # Extract the response content
        response_content = completion.choices[0].message.content

        # Check if the response content is in JSON format and strip if needed
        if response_content.startswith("```json"):
            response_content = response_content.strip("```json\n").strip("```")
        
        # Attempt to parse response content as JSON
        try:
            return json.loads(response_content)
        except json.JSONDecodeError as json_error:
            logging.error(f"Failed to decode JSON response: {json_error}")
            logging.error(f"Response content was: {response_content}")
            return None
    
    except Exception as e:
        logging.warning(f"An error occurred: {e}")
        return None

def Query_OpenAI_Assistant(assistant_id, query_text):
    # Call the OpenAI API for a chat completion from a predefined assistant
    myassistant = OpenAIclient.beta.assistants.retrieve(assistant_id)
    #print(myassistant.name)

    mythread = OpenAIclient.beta.threads.create()
    #print(empty_thread)

    message = OpenAIclient.beta.threads.messages.create(
        thread_id=mythread.id,
        role="user",
        content=query_text
    )

    run = OpenAIclient.beta.threads.runs.create_and_poll(
        thread_id=mythread.id,
        assistant_id=myassistant.id
    )

    response = ""
    if run.status == 'completed': 
        messages = OpenAIclient.beta.threads.messages.list(
            thread_id=mythread.id
        )
        response = messages.data[0].content[0].text.value
    else:
        print(run.status)

    return response

def extract_text_from_docx(docx_path):
    """
    Extracts all text from a .docx file and returns it as a string.
    """
    try:
        doc = Document(docx_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        logging.error(f"Error reading DOCX file {docx_path}: {e}")
        return ""

def anonymizeName(file_body, consultant_name, consultant_alias):
    # Split the name field by space
    name_split = consultant_name.split()
    num_names = len(name_split)
    # Assign the first element to first_name
    first_name = name_split[0]
    # Assign the last element to last_name
    last_name = name_split[num_names - 1]
    # Place the middle name correctly
    if num_names > 2:
        middle_name = name_split[1]
        if middle_name in ('Holte', 'Sogn', 'Ølfarnes'):
            last_name = middle_name + " " + last_name
        else:
            if len(middle_name) > 1: # Ignore abbreviated middle names
                first_name += " " + middle_name

    # Define the first_name and last_name
    short_name = name_split[0]
    full_name  = first_name + " " + last_name
    genetiv    = first_name + "s"

    # Anonymize every row; Replace first name by "Konsulenten(s)"
    content = file_body
    content = re.sub(r"\b%s\b" %  consultant_name,  consultant_alias, content)
    content = re.sub(r"\b%s\b" %  full_name,  consultant_alias, content)
    content = re.sub(r"\b%s\b" %  genetiv,    consultant_alias+"s", content)
    content = re.sub(r"\b%s\b" %  first_name, consultant_alias, content)
    content = re.sub(r"\b%s\b" %  last_name,  consultant_alias, content)
    content = re.sub(r"\b%s\b" %  short_name, consultant_alias, content)

    return content

def extract_competency_matrix(df_consultants, cv_folder="downloaded_CVs"):
    """
    Creates a competency matrix by analyzing each consultant's CV and extracting competency details.
    
    Args:
        df (DataFrame): DataFrame containing consultant data with 'Ressurs' (consultant name) column.
        cv_folder (str): Directory where downloaded CVs are stored.
    
    Returns:
        DataFrame: Competency matrix with one row per consultant.
    """
    competency_data = []
    competency_assistant_id = "asst_ommyrTiCpM234i21Nqj25MkW"

    # Ensure profiles folder exists
    os.makedirs(PROFILES_FOLDER, exist_ok=True)

    # Make column explanation table
    all_themes = {}
    for schemafile in CompetencyBotSchemas:
        # Read schema-file
        if os.path.exists(schemafile):
            with open(schemafile, "r", encoding="utf-8") as f:
                schema_text = f.read()
                #logging.info(f"Schema file loaded: ({schemafile}).")
        else:
            logging.error(f"No schemafile {CONSULTANT_JSON_FILE} found!  Did you forget to run GenererSchema.py first?")
            exit(1)

        try:
            schema_json = json.loads(schema_text)
            # Logging or print to confirm successful parsing
            logging.info(f"Schema JSON successfully parsed: {schemafile}")
        except json.JSONDecodeError as e:
            logging.error(f"Error parsing JSON from schema file {schemafile}: {e}")
            exit(1)

        all_themes.update(schema_json)

    # Make row that explains variables
    comp_explanation = {}
    for theme in all_themes.keys():
        if theme.endswith(".ref"):
            comp_explanation[theme] = all_themes[theme].replace("Kort referanse. ", "")
        else:
            comp_explanation[theme] = ""
    competency_data.append({
        "Consultant": "FORKLARING",
        "Competency": comp_explanation
    })

    # Iterate over each consultant in the DataFrame and ask assistant to create a row in competency matrix
    for _, row in df_consultants.iterrows():
        consultant_name = row['Ressurs']
        cv_path = os.path.join(cv_folder, f"{consultant_name}_CV.docx")
        profile_path = os.path.join(PROFILES_FOLDER, f"{consultant_name}_profile.json")
        
        if os.path.exists(cv_path):
            # Get current CV's modification timestamp
            current_timestamp = int(os.path.getmtime(cv_path))
            
            # Check if a valid profile already exists
            if os.path.exists(profile_path):
                with open(profile_path, 'r', encoding="utf-8") as profile_file:
                    profile_data = json.load(profile_file)
                    last_filename = profile_data.get("CV filnavn")
                    last_timestamp = profile_data.get("CV timestamp")

                    # Skip profiling if the current CV matches the last analyzed CV
                    if last_filename == f"{consultant_name}_CV.docx" and last_timestamp == current_timestamp:
                        logging.info(f"Profile for {consultant_name} is up-to-date. Skipping re-profiling.")
                        competency_data.append({
                            "Consultant": consultant_name,
                            "Competency": profile_data["Competency"]
                        })
                        continue

            # Convert CV to text and anonymize it by removing consultant names
            logging.info(f"Extracting text from {consultant_name}'s CV.")
            cv_text = extract_text_from_docx(cv_path)
            anonymized_cv = anonymizeName(cv_text, consultant_name, "Konsulenten")
            Full_competency_info = {}

            for schemafile in CompetencyBotSchemas:
                # Read schema-file
                if os.path.exists(schemafile):
                    with open(schemafile, "r", encoding="utf-8") as f:
                        schema_text = f.read()
                        #logging.info(f"Schema file loaded: ({schemafile}).")
                else:
                    logging.error(f"No schemafile {CONSULTANT_JSON_FILE} found!  Did you forget to run GenererSchema.py first?")
                    exit(1)

                #Build full instruction
                CompetencyBotInstruction = f"{CompetencyBotIntro}{schema_text}\n{CompetencyBotOutro}"

                # Analyze text with OpenAI assistant
                try:
                    logging.info(f"Analyzing competency for {consultant_name} with schema {schemafile}.")
                    competency_info = Query_OpenAI_Completion_JSON(CompetencyBotInstruction, anonymized_cv)
                    # logging.info(f"Competency extracted for {consultant_name}: {competency_info}")
                    Full_competency_info.update(competency_info)
                except Exception as e:
                    logging.error(f"Error analyzing competency for {consultant_name} with schema {schemafile}: {e}")
                    Full_competency_info.update({
                        "SYSTEM ERROR DURING AI EXTRACTION": f"Error for {schemafile}: {e}"
                    })

            # Append results to the data list
            competency_data.append({
                "Consultant": consultant_name,
                "Competency": Full_competency_info
            })

            # Save individual profile to JSON
            profile_data = {
                "Consultant": consultant_name,
                "CV filnavn": f"{consultant_name}_CV.docx",
                "CV timestamp": current_timestamp,
                "Competency": Full_competency_info
            }
            with open(profile_path, 'w', encoding="utf-8") as profile_file:
                json.dump(profile_data, profile_file, indent=4)
            logging.info(f"Profile saved for {consultant_name}.")
                
        else:
            logging.warning(f"CV file not found for {consultant_name}.")
            competency_data.append({
                "Consultant": consultant_name,
                "Competency": "CV file not found"
            })

    # Convert the competency data to a DataFrame
    competency_df = pd.DataFrame(competency_data)
    logging.info("Competency matrix extraction completed.")
    return competency_df

def profil_kvalitetssikring():
    folder_path = "consultant_profiles"
    files = [f for f in os.listdir(folder_path) if f.endswith('.json')]
    
    if not files:
        print("Ingen JSON-filer funnet i mappen.")
        return
    
    # Les inn den første filen som referanse
    with open(os.path.join(folder_path, files[0]), 'r', encoding='utf-8') as f:
        first_file_data = json.load(f)
        competency_keys = set(first_file_data.get("Competency", {}).keys())

    # Sjekk de andre filene i mappen
    for file in files[1:]:
        with open(os.path.join(folder_path, file), 'r', encoding='utf-8') as f:
            data = json.load(f)
            other_competency_keys = set(data.get("Competency", {}).keys())

            # Sjekk for manglende og ekstra nøkler
            missing_keys = competency_keys - other_competency_keys
            extra_keys = other_competency_keys - competency_keys

            if missing_keys:
                print(f"Warning! File {file} lacks fields: {', '.join(missing_keys)}.")
            if extra_keys:
                print(f"Warning! File {file} has additional fields: {', '.join(extra_keys)}.")

def expand_competency_data(df):
    # Assuming the JSON data is in a column called 'Competency'
    df['Competency'] = df['Competency'].apply(lambda x: json.loads(x) if isinstance(x, str) else x)
    competency_expanded_df = pd.json_normalize(df['Competency'])
    df = df.drop(columns=['Competency']).join(competency_expanded_df)
    return df

def format_norwegian_locale(x):
    if isinstance(x, float):
        return f"{x:.2f}".replace('.', ',')
    return x

DEBUG = False

# Main execution
if __name__ == "__main__":
    if DEBUG:
        logging.warning(" *** DEBUGGING MODE ***")

    logging.info("Loading local consultant data file:")
    # Load JSON data if file exists
    consultant_data = {}
    if os.path.exists(CONSULTANT_JSON_FILE):
        with open(CONSULTANT_JSON_FILE, "r", encoding="utf-8") as f:
            consultant_data = json.load(f)
        logging.info(f"Consultant CV data loaded from existing JSON file ({CONSULTANT_JSON_FILE}).")
    else:
        logging.error(f"No Consultant Data file found ({CONSULTANT_JSON_FILE}). Did you forget to run Step1 to completion?")
        exit(1)

    logging.info("Loading local copy of A-2 PIPELINE:")
    if os.path.exists("a2pipeline.csv"):
        pipeline_df = pd.read_csv("a2pipeline.csv")
        logging.info(f"Pipelinedata loaded from existing CSV file.")
    else:
        logging.error(f"No A2 Pipeline Data file found (a2pipeline.csv). Did you forget to run Step1 to completion?")
        exit(1)

    if DEBUG:
        pattern = "Eirik|Henning|Ingrid|Karl|Carol|Desir|Molan"
        pipeline_df = pipeline_df[pipeline_df['Ressurs'].str.contains(pattern, na=False, regex=True)]
        logging.warning(" *** DEBUGGING MODE ***")
        #print(pipeline_df)

    # Step 1: Extract competency matrix
    logging.info("EXTRACTING COMPETENCY MATRIX:")
    competency_matrix_df = extract_competency_matrix(pipeline_df)

    # Run QA on all profiles
    profil_kvalitetssikring()

    # Step 2: Expand JSON fields into individual columns
    competency_matrix_df = expand_competency_data(competency_matrix_df)
    #print("Column names after expansion:", competency_matrix_df.columns)

    # Step 3: Ensure correct encoding and save to CSV
    output_file = "competency_matrix_cleaned.csv"

    # Replace NaN values with empty strings
    competency_matrix_df = competency_matrix_df.fillna("")

    # Format the DataFrame for Norwegian locale settings using map
    competency_matrix_df = competency_matrix_df.map(format_norwegian_locale)

    # Save the DataFrame to a CSV file with a semicolon delimiter
    competency_matrix_df.to_csv(output_file, sep=';', index=False, encoding='utf-8-sig')
    logging.info(f"Competency matrix saved to {output_file} with UTF-8 encoding.")